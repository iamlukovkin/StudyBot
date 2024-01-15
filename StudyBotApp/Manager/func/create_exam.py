import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

from database import *
from Manager.util import my_docx


def _(filename: str, param: str, search_by: str = 'group'):
    """
    Creates docx document.

    Args:
        group (str): group name
    """
    
    doc = my_docx.get_docx()
    h1 = doc.add_paragraph(filename, style='Heading 1')
    h1.style.color = RGBColor(255, 0, 128)
    h1.style.font.name = 'Arial'
    h1.style.font.size = Pt(32)
    h1.style.bold = True
    
    if search_by == 'group':
        sessions = get_session().query(Session).filter(
            Session.group == param
        ).order_by(Session.date).all()
        show_group = False
    
    else:
        sessions = get_session().query(Session).filter(
            Session.exam_info.like('%{}%'.format(param))
        ).order_by(Session.id).all()
        show_group = True
        
    for exam in sessions:
        exam: Session
        if show_group:
            group = exam.group.split('_')[1]
            p = doc.add_paragraph('Группа: ' + group, style='Heading 3')
            p.style.font.name = 'Arial'
            p.style.font.size = Pt(18)
            p.style.bold = True
        p = doc.add_paragraph(exam.exam_info)
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(18)
        p.style.bold = True
        p = doc.add_paragraph(
            'Дата: {}'.format(exam.date) + ' время: {}\n\n'.format(
                    exam.time.strftime("%H:%M")),
        )
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(12)
    
    return doc
        
        
        
        
    
    
        
        