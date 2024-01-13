import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

from database import *
from Manager.util import my_docx


def _(filename: str, param: str, search_by: str = 'group'):
    """
    Creates docx document.

    Args:
        group (str): group name
    """
    
    doc = my_docx.get_docx()
    h1 = doc.add_paragraph(filename, style='Heading 1')
    h1.style.color = RGBColor(255, 0, 255)
    h1.style.font.name = 'Arial'
    h1.style.font.size = Pt(32)
    h1.style.bold = True
    
    weeks = ['Числитель', 'Знаменатель']
    days = ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота']
    
    for week in weeks:
        p = doc.add_paragraph(week, style='Heading 2')
        p.style.color = RGBColor(255, 0, 255)
        p.style.font.name = 'Arial'
        p.style.font.size = Pt(24)
        p.style.bold = True
        for day in days:
            session = get_session()
            if search_by == 'group':
                lessons = session.query(
                    Lesson
                    ).filter(Lesson.group == param
                    ).filter(Lesson.day == day
                    ).filter(Lesson.week == week
                    ).order_by(Lesson.time_start
                ).all()
            
            else:
                lessons = session.query(
                    Lesson
                    ).filter(Lesson.lesson_info.like('%{}%'.format(param))
                    ).filter(Lesson.day == day
                    ).filter(Lesson.week == week
                    ).order_by(Lesson.time_start
                ).all()
            
            if len(lessons) == 0:
                continue
            p = doc.add_paragraph(day, style='Heading 3')
            p.style.color = RGBColor(255, 0, 255)
            p.style.font.name = 'Arial'
            p.style.font.size = Pt(18)
            p.style.bold = True
            
            later_time = lessons[0].time_start
            groups: list = []    
            
            day_dict: dict = {}
            for lesson in lessons:
                lesson: Lesson
                
                string_text: str = '{0}\n' 
                string_text += 'Начало: {1}' 
                string_text += '\nПерерыв: {2}'
                
                if lesson.time_end:
                    string_text += '\nКонец: {3}' 
                
                if lesson.time_start not in day_dict:
                    start_time, break_time, end_time = lesson.get_times()
                    day_dict[lesson.time_start] = [
                        [lesson.get_group()], 
                        string_text.format(
                            lesson.lesson_info,
                            start_time.strftime('%H:%M'),
                            break_time.strftime('%H:%M'),
                            end_time.strftime('%H:%M')
                        )
                    ]
                else:
                    day_dict[lesson.time_start][0].append(
                        lesson.get_group())
            
            for time in day_dict.keys():
                if len(day_dict[time][0]) == 1:
                    prefix = 'Группа: '
                else:
                    prefix = 'Группы: '
                groups: str = prefix + ', '.join(day_dict[time][0]) 
                p = doc.add_paragraph(groups)
                p.style.color = RGBColor(0, 0, 0)
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(12)
                p.style.bold = True
                p = doc.add_paragraph(day_dict[time][1] + '\n\n')
                p.style.color = RGBColor(0, 0, 0)
                p.style.font.name = 'Arial'
                p.style.font.size = Pt(12)
    
    return doc
